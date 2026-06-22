"""Build a Word document of the June 2026 newsletter with image slots.
Local artifact only — does NOT touch the brain or Google Drive."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

ACCENT = RGBColor(0xB5, 0x49, 0x1F)
GREY = RGBColor(0x55, 0x55, 0x55)
OUT = os.path.join(os.path.dirname(__file__), "newsletter-jun2026.docx")

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)

def body(text, italic=False, size=11, space_after=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def headline(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def image_slot(label):
    """A bordered, shaded placeholder cell — drag your image in here, then delete the text."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    cell.width = Inches(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(f"[ IMAGE — {label} ]")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("• • •")
    r.font.color.rgb = GREY

# ---------- Opening ----------
body("Dear <<First Name>>,")
body("Hope this message finds you in good spirits.")
body("There's a quiet thread running through this edition: a movement rooted in one place rarely "
     "stays there — it travels. A film born in the villages of Bihar finds its way to Los Angeles "
     "and the ABBY stage. The same question asked on the streets of Anekal — does every child have "
     "a school worth believing in? — is the one carried into a room in San Francisco and a long "
     "table at Oxford. What holds it together is ownership: parents, teachers, youth, communities, "
     "partners, and the diaspora, each beginning to see themselves as co-owners of public education's "
     "future.")
body("This is what Shikshagraha looks like in action — across chaupals, streets, and convenings, "
     "here and across the world. This edition is yet another window into the movement, in practice.")
body("Happy reading...", italic=True)
divider()

# ---------- Block 1: Mothers of Courage ----------
image_slot("Mothers of Courage")
headline("From Bihar's Villages to the World's Stages: Mothers of Courage")
body("Mothers of Courage — the film following Bihar's Shiksha Chaupals — has won Silver in the "
     "SDG category at the ABBY Awards 2026, adding to Best Inspirational Film at the Los Angeles Film "
     "Awards and a feature at the Cannes Lions ERA Programme. The News Minute went behind the film: a "
     "crew that lived alongside these women until the camera all but disappeared and the real "
     "conversations began. 28,000+ Chaupals. 14 districts. 8 lakh+ people moved to act.")
body("Read The News Minute coverage: https://lnkd.in/gnMtBJkg", size=10, color=ACCENT, space_after=2)
body("Support the mothers leading this change — donate: https://bit.ly/49eRUgk", size=10, color=ACCENT)
divider()

# ---------- Block 2: Anekal ----------
image_slot("Anekal street march")
headline("Every Child to School: Anekal Takes the Message to the Streets")
body("On 29 May, over 350 people — government officials, NGOs, teachers, youth leaders, and "
     "community members — took to the streets of Anekal for every child's right to free, equal, "
     "quality education. Youth led the way: 55+ young leaders made placards, hoardings, and recorded "
     "enrolment messages. Slogans like “Every Child to School” and “Government Schools, "
     "Our Pride” spread through local media and community networks well beyond the day.")
body("[CTA link — to be shared]", size=10, color=ACCENT)
divider()

# ---------- Block 3: Glocal ----------
image_slot("Unbundling Evaluation forum")
headline("Unbundling Evaluation: Communities at the Centre of How We Learn")
body("As part of GLOCAL Evaluation Week, the Centre for Exponential Change, ShikshaLokam, and CoLab "
     "hosted “Unbundling Evaluation” — a “one-size-fits-one” forum on closing the "
     "gap between evaluation theory and practice. Three insights stood out: sensemaking is the most "
     "under-resourced step in qualitative research; feedback loops are rarely closed with the "
     "communities who provide the data; and impact measurement too often misses the relational "
     "milestones that matter most — a first referral, a gained sense of identity.")
body("[CTA link — to be shared]", size=10, color=ACCENT)
divider()

# ---------- Block 4: MItra ----------
image_slot("MItra")
headline("Stories That Matter: A Tool That Listens Like a Friend")
body("India's 1.5 lakh schools sit on a goldmine of solutions — but they stay locked where they "
     "were born. The teacher in Nagaland never hears the attendance trick from Karnataka; the mother in "
     "Bihar doesn't know hundreds of others are fighting the same battle elsewhere. So we built MItra "
     "— a voice-enabled tool that lets school, community, and youth leaders share their "
     "micro-improvements in their own language, turning quiet conversations into stories worth learning "
     "from. An education leader's friend.")
body("First in our three-part series, Stories That Matter.", italic=True)
body("[CTA link — to be shared]", size=10, color=ACCENT)
divider()

# ---------- Block 5: Visits ----------
image_slot("Portola Valley & Skoll")
headline("From Portola Valley to Oxford: Shikshagraha in Global Conversation")
body("In Portola Valley, San Francisco, philanthropists, entrepreneurs, tech leaders, and members of "
     "the Indian diaspora gathered around one question: what will it take for every child in India's "
     "one million schools to be future-ready? The room invited the diaspora in as active partners, not "
     "distant supporters — with Mantra4Change US building bridges of collective action across "
     "geographies.")
body("Weeks later, at the Skoll World Forum in Oxford, Shikshagraha co-convened Education House with "
     "Teach For All, Salzburg Global, and Global Schools Forum (200+ participants), alongside a research "
     "roundtable with the Government Outcomes Lab and a listening circle with the Schwab Foundation and "
     "Synergos. The week opened doors to 30+ new partners.")
body("[CTA link — to be shared]", size=10, color=ACCENT)
divider()

# ---------- Footer ----------
body("Stay updated on the latest stories in education leadership. Follow us!", italic=True, color=GREY, space_after=4)
body("LinkedIn · YouTube · Website · Spotify · Twitter", size=10, space_after=4)
body("Write to us with your feedback at social@shikshalokam.org", size=10, color=GREY, space_after=4)
body("To unsubscribe, reply with a 'U' to this email.", size=9, color=GREY, space_after=14)
body("Warm regards,", space_after=2)
p = doc.add_paragraph(); r = p.add_run("Khushboo Awasthi"); r.bold = True
body("Chief Curator, Shikshagraha", size=10, color=GREY)

doc.save(OUT)
print("Wrote", OUT)
