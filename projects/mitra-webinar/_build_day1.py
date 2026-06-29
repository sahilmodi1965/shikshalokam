#!/usr/bin/env python3
"""Build a Word doc of the MItra webinar Day-1 assets and report the path."""
import re
from pathlib import Path
from docx import Document

REPO = Path(__file__).resolve().parents[2]
page = (REPO / "projects" / "mitra-webinar" / "page.md").read_text(encoding="utf-8")

# everything in the asset library, keep only Day 1 blocks
lib = page.split("## Asset library", 1)[1]
blocks = ("\n" + lib).split("\n### ")
day1 = [b for b in blocks if b.startswith("Day 1")]
body = "### " + "\n### ".join(day1)

doc = Document()
doc.add_heading("MItra Webinar — Day 1 Assets", level=0)

def add_runs(para, text):
    text = text.replace("`", "")
    for t in re.split(r'(\*\*.+?\*\*|\*[^*]+?\*)', text):
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            para.add_run(t[2:-2]).bold = True
        elif t.startswith("*") and t.endswith("*"):
            para.add_run(t[1:-1]).italic = True
        else:
            para.add_run(t)

def styled(name):
    try:
        doc.add_paragraph(style=name)
        return doc.paragraphs[-1]
    except KeyError:
        return doc.add_paragraph()

for raw in body.splitlines():
    line = raw.rstrip()
    if not line.strip():
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=1)
    elif line.startswith("> "):
        add_runs(styled("Quote"), line[2:])
    elif line.strip() == ">":
        continue
    elif line.startswith("- "):
        add_runs(doc.add_paragraph(style="List Bullet"), line[2:])
    else:
        add_runs(doc.add_paragraph(), line)

out = REPO / "projects" / "mitra-webinar" / "MItra Webinar - Day 1 Assets.docx"
doc.save(str(out))
print("saved:", out)
